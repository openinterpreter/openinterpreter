"""
Document parser for LegalDesk.

Extracts text from PDF, DOCX, and plain text files for legal analysis.
Supports OCR fallback for scanned PDFs via Tesseract.
"""

import os


def extract_text(file_path):
    """
    Extract text from a document file.

    Supports:
    - PDF (via PyPDF2, with OCR fallback via pytesseract)
    - DOCX (via python-docx)
    - TXT / MD (plain text)

    Returns the extracted text as a string.
    Raises ValueError if the file type is not supported.
    """
    ext = os.path.splitext(file_path)[1].lower()

    if ext == ".pdf":
        return extract_from_pdf(file_path)
    elif ext == ".docx":
        return extract_from_docx(file_path)
    elif ext in (".txt", ".md", ".text", ".rst"):
        return extract_from_text(file_path)
    else:
        raise ValueError(
            f"Unsupported file type: {ext}. Supported types: .pdf, .docx, .txt, .md"
        )


def extract_from_pdf(file_path):
    """Extract text from a PDF file. Falls back to OCR if text extraction yields little content."""
    try:
        import PyPDF2
    except ImportError:
        raise ImportError(
            "PyPDF2 is required for PDF parsing. Install it with: pip install PyPDF2"
        )

    text_parts = []
    with open(file_path, "rb") as f:
        reader = PyPDF2.PdfReader(f)
        for page in reader.pages:
            page_text = page.extract_text()
            if page_text:
                text_parts.append(page_text)

    text = "\n\n".join(text_parts)

    # If very little text was extracted, try OCR
    if len(text.strip()) < 100:
        ocr_text = _try_ocr_pdf(file_path)
        if ocr_text and len(ocr_text.strip()) > len(text.strip()):
            return ocr_text

    return text


def _try_ocr_pdf(file_path):
    """Attempt OCR on a PDF file using pytesseract. Returns None if dependencies missing."""
    try:
        import pytesseract
        from pdf2image import convert_from_path
    except ImportError:
        return None

    try:
        images = convert_from_path(file_path)
        text_parts = []
        for image in images:
            text_parts.append(pytesseract.image_to_string(image))
        return "\n\n".join(text_parts)
    except Exception:
        return None


def extract_from_docx(file_path):
    """Extract text from a DOCX file."""
    try:
        import docx
    except ImportError:
        raise ImportError(
            "python-docx is required for DOCX parsing. Install it with: pip install python-docx"
        )

    doc = docx.Document(file_path)
    paragraphs = [para.text for para in doc.paragraphs if para.text.strip()]
    return "\n\n".join(paragraphs)


def extract_from_text(file_path):
    """Extract text from a plain text file."""
    with open(file_path, "r", encoding="utf-8") as f:
        return f.read()


def get_document_metadata(file_path):
    """Get basic metadata about a document."""
    ext = os.path.splitext(file_path)[1].lower()
    stat = os.stat(file_path)

    metadata = {
        "filename": os.path.basename(file_path),
        "extension": ext,
        "size_bytes": stat.st_size,
        "size_human": _human_readable_size(stat.st_size),
    }

    if ext == ".pdf":
        try:
            import PyPDF2

            with open(file_path, "rb") as f:
                reader = PyPDF2.PdfReader(f)
                metadata["page_count"] = len(reader.pages)
                if reader.metadata:
                    metadata["title"] = reader.metadata.get("/Title", "")
                    metadata["author"] = reader.metadata.get("/Author", "")
        except Exception:
            pass

    elif ext == ".docx":
        try:
            import docx

            doc = docx.Document(file_path)
            metadata["paragraph_count"] = len(doc.paragraphs)
            core = doc.core_properties
            metadata["title"] = core.title or ""
            metadata["author"] = core.author or ""
        except Exception:
            pass

    return metadata


def _human_readable_size(size_bytes):
    """Convert bytes to human-readable string."""
    for unit in ["B", "KB", "MB", "GB"]:
        if size_bytes < 1024:
            return f"{size_bytes:.1f} {unit}"
        size_bytes /= 1024
    return f"{size_bytes:.1f} TB"
